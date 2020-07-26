#! /usr/bin/env python3
# Playing with DOCXs
# Official documentation for python-docx -> https://python-docx.readthedocs.io/en/latest/
import platform
import pathlib
import subprocess
from distutils.spawn import find_executable

import docx

demo_file = pathlib.Path("input_docx/demo.docx")
print("Playing with DOCXs")
print(f"Reading a doc file: {demo_file.absolute()}")
doc = docx.Document(demo_file)
print(f"Number of paragraphs: {len(doc.paragraphs)}")
for p in doc.paragraphs:
    print(p.text)
    print(f"{'*' * 30}")

hello_world_file = pathlib.Path("output_docx/helloworld.docx")
doc_hello_word = docx.Document()
doc_hello_word.add_heading("Hello Title", 0)
p = doc_hello_word.add_paragraph("Hello, ")
p.add_run(" world!").bold = True
doc_hello_word.add_heading("Programming Languages", level=1)
for language in ["C", "Java", "Python", "Ruby"]:
    doc_hello_word.add_paragraph(language, style="List Bullet")
doc_hello_word.add_page_break()
doc_hello_word.add_paragraph("Other info")
doc_hello_word.save(hello_world_file)
print(f"Writing a doc file: {hello_world_file.absolute()}")

if platform.system() == "Linux":
    soffice_path = find_executable("soffice")
    if soffice_path:
        print("Generating a PDF file using LibreOffice command line utility")
        cmd = [soffice_path, "--convert-to", "pdf", hello_world_file.absolute(), "--outdir", "output_pdf"]
        subprocess.call(cmd)
